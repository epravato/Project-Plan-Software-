"""
Builds the H2M project plan Word document from scratch (no template file).

Generates a clean, report-style document rather than filling H2M's old rigid
form-style template — narrative sections read as prose, lookups (contacts,
team, milestones, risks) stay in clean tables. Section order/numbering matches
H2M's official "Minimum Requirements of a Project Plan" (12 required sections).
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

BLUE = RGBColor(0x67, 0x90, 0xB1)
BLUE_DARK = RGBColor(0x4A, 0x75, 0x94)
GREEN = RGBColor(0xA2, 0xC0, 0x37)
INK = RGBColor(0x2E, 0x2E, 0x30)
MUTED = RGBColor(0x6B, 0x6B, 0x6E)
BORDER_HEX = "D3DAE1"

NO_VALUE = "Not yet provided."
NO_ROWS = "None entered."


# ------------------------------------------------------------------ oxml helpers

def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bottom_border(paragraph, hex_color, size=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


def _tile(cell, letter, fill):
    """One square of the H2M logo mark: white letter on a brand-color tile."""
    cell.text = ""
    cell.width = Inches(0.26)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if letter:
        run = p.add_run(letter)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if fill:
        _shade(cell, fill)


def _shrink_row_height(row, twips=140):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(twips))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


# ------------------------------------------------------------------ styles

def _setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)

    title = doc.styles["Title"]
    title.font.name = "Georgia"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(2)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Georgia"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = BLUE_DARK
    h1.paragraph_format.space_before = Pt(22)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(11.5)
    h2.font.bold = True
    h2.font.color.rgb = INK
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True


def _header_footer(doc, project_label):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ""

    # H2M logo mark: 2x2 tile grid (H / 2 / blank / M) plus the tagline, built
    # from shaded cells rather than an image so there's no asset dependency.
    logo = header.add_table(rows=2, cols=3, width=Inches(6.7))
    _no_borders(logo)
    _tile(logo.cell(0, 0), "H", "A2C037")
    _tile(logo.cell(0, 1), "2", "58585A")
    _tile(logo.cell(1, 0), "", None)
    _tile(logo.cell(1, 1), "M", "6790B1")

    name_cell = logo.cell(0, 2)
    name_cell.text = ""
    name_cell.width = Inches(6.0)
    np = name_cell.paragraphs[0]
    np.paragraph_format.space_after = Pt(0)
    np.paragraph_format.left_indent = Inches(0.12)
    nr = np.add_run("architects + engineers")
    nr.font.name = "Calibri Light"
    nr.font.size = Pt(13)
    nr.font.color.rgb = RGBColor(0x8C, 0x8C, 0x8E)

    tag_cell = logo.cell(1, 2)
    tag_cell.text = ""
    tp = tag_cell.paragraphs[0]
    tp.paragraph_format.space_after = Pt(0)
    tp.paragraph_format.left_indent = Inches(0.12)
    t1 = tp.add_run("practical approach. ")
    t1.font.name = "Calibri Light"
    t1.font.size = Pt(8)
    t1.font.color.rgb = RGBColor(0x8C, 0x8C, 0x8E)
    t2 = tp.add_run("creative results.")
    t2.font.name = "Calibri Light"
    t2.font.size = Pt(8)
    t2.font.color.rgb = RGBColor(0xC0, 0xA1, 0x6B)

    rule = header.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(0)
    _bottom_border(rule, "A2C037", size=12)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    frun = fp.add_run(project_label or "H2M Project Plan")
    frun.font.size = Pt(8)
    frun.font.color.rgb = MUTED
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ------------------------------------------------------------------ content helpers

def _heading(doc, number, title):
    p = doc.add_heading(level=1)
    p.text = ""
    run = p.add_run(f"{number}.  {title.upper()}")
    run.font.name = "Georgia"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK
    _bottom_border(p, BORDER_HEX)
    return p


def _subheading(doc, label, title):
    p = doc.add_heading(level=2)
    p.text = ""
    run = p.add_run(f"{label}. {title}")
    run.font.color.rgb = INK
    return p


def _body(doc, text):
    p = doc.add_paragraph()
    if text:
        p.add_run(text)
    else:
        run = p.add_run(NO_VALUE)
        run.italic = True
        run.font.color.rgb = MUTED
    return p


def _meta_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = MUTED
    r2 = p.add_run(value or "TBD")
    r2.font.size = Pt(9.5)
    return p


def _table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    head = table.rows[0]
    for i, text in enumerate(headers):
        cell = head.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text.upper())
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, "6790B1")
    _shrink_row_height(head)

    if not rows:
        row = table.add_row()
        cell = row.cells[0]
        for c in row.cells[1:]:
            cell.merge(c)
        run = cell.paragraphs[0].add_run(NO_ROWS)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = MUTED
    else:
        for values in rows:
            row = table.add_row()
            for i, value in enumerate(values):
                cell = row.cells[i]
                run = cell.paragraphs[0].add_run(str(value) if value else "")
                run.font.size = Pt(9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            if w is not None:
                for row in table.rows:
                    row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ------------------------------------------------------------------------ build

def build(
    output_path: str | Path,
    fields: dict[str, str] | None = None,
    team: list[dict] | None = None,
    client_success: list[dict] | None = None,
    h2m_success: list[dict] | None = None,
    meetings: list[dict] | None = None,
    drawings: list[dict] | None = None,
    specifications: list[dict] | None = None,
    cost_opinions: list[dict] | None = None,
    reports: list[dict] | None = None,
    milestones: list[dict] | None = None,
    risks: list[dict] | None = None,
    statement_of_purpose: str | None = None,
    hazard_assessment: str | None = None,
    scope_of_work: str | None = None,
    scope_of_services: str | None = None,
    qa_qc_plan: str | None = None,
    wbs_link: str | None = None,
    schedule_link: str | None = None,
    bim_link: str | None = None,
    fee_link: str | None = None,
    invoice_frequency: str | None = None,
    invoice_date: str | None = None,
    progress_frequency: str | None = None,
    progress_format: str | None = None,
    progress_delivery: str | None = None,
):
    fields = fields or {}
    doc = Document()
    _setup_styles(doc)

    project_name = fields.get("Project Name") or "Untitled Project"
    client_name = fields.get("Client Company") or ""
    _header_footer(doc, f"{project_name} — Project Plan")

    # ---- Title block
    doc.add_paragraph().text = ""
    title = doc.add_paragraph(style="Title")
    title.add_run(project_name)
    if client_name:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(16)
        run = sub.add_run(client_name)
        run.italic = True
        run.font.size = Pt(13)
        run.font.color.rgb = MUTED

    # ---- I. General Project Information
    _heading(doc, "I", "General Project Information")
    _table(
        doc,
        ["", "H2M Project Team", "Client"],
        [
            ["Contact", fields.get("Project Contact", ""), fields.get("Client Contact", "")],
            ["Location", fields.get("Project Location", ""), fields.get("Client Location", "")],
            ["Phone", fields.get("Project Phone", ""), fields.get("Client Phone", "")],
            ["Fax", fields.get("Project Fax", ""), fields.get("Client Fax", "")],
            ["E-Mail", fields.get("Project E-Mail", ""), fields.get("Client E-Mail", "")],
        ],
        col_widths=[1.1, 2.6, 2.6],
    )

    # ---- II. Project Team
    _heading(doc, "II", "Project Team")
    _table(
        doc, ["Name", "Organization", "Role", "Email", "Phone"],
        [[m.get("name", ""), m.get("organization", ""), m.get("role", ""),
          m.get("email", ""), m.get("phone", "")] for m in (team or [])],
    )

    # ---- III. Overview
    _heading(doc, "III", "Overview")
    _subheading(doc, "A", "Statement of Purpose")
    _body(doc, statement_of_purpose)
    _subheading(doc, "B", "Critical Success Factors")
    p = doc.add_paragraph()
    r = p.add_run("Client's Critical Success Factors")
    r.bold = True
    r.font.size = Pt(10)
    _table(
        doc, ["Success Factor", "Performance Objective"],
        [[f.get("factor", ""), f.get("metric", "")] for f in (client_success or [])],
    )
    p = doc.add_paragraph()
    r = p.add_run("H2M's Critical Success Factors")
    r.bold = True
    r.font.size = Pt(10)
    _table(
        doc, ["Success Factor", "Performance Objective"],
        [[f.get("factor", ""), f.get("metric", "")] for f in (h2m_success or [])],
    )

    # ---- IV. Health & Safety
    _heading(doc, "IV", "Health & Safety")
    _subheading(doc, "A", "Project Hazard Assessment")
    _body(doc, hazard_assessment or "To be completed — refer to the H2M Project Hazard Assessment Form.")

    # ---- V. Project Deliverables
    _heading(doc, "V", "Project Deliverables")
    _subheading(doc, "A", "Meetings")
    _table(
        doc, ["Type", "Frequency", "Anticipated Attendees (By Title)"],
        [[m.get("type", ""), m.get("frequency", ""), m.get("attendees", "")] for m in (meetings or [])],
    )
    _subheading(doc, "B", "Drawings")
    _table(
        doc, ["Discipline", "Dwg. #", "Dwg. Title"],
        [[d.get("discipline", ""), d.get("number", ""), d.get("title", "")] for d in (drawings or [])],
    )
    _subheading(doc, "C", "Specifications")
    _table(
        doc, ["Div. - Section No.", "Specification Section Name"],
        [[s.get("section_no", ""), s.get("name", "")] for s in (specifications or [])],
    )
    _subheading(doc, "D", "Cost Opinions")
    _table(
        doc, ["Milestone", "Type / Title", "Link / Comments"],
        [[r_.get("milestone", ""), r_.get("title", ""), r_.get("comments", "")] for r_ in (cost_opinions or [])],
    )
    _subheading(doc, "E", "Reports")
    _table(
        doc, ["Milestone", "Type / Title", "Link / Comments"],
        [[r_.get("milestone", ""), r_.get("title", ""), r_.get("comments", "")] for r_ in (reports or [])],
    )

    # ---- VI. Work Breakdown Structure
    _heading(doc, "VI", "Work Breakdown Structure")
    _meta_line(doc, "Link to detailed WBS", wbs_link)

    # ---- VII. Scope of Work & Services
    _heading(doc, "VII", "Scope of Work & Services")
    _subheading(doc, "A", "Scope of Work")
    _body(doc, scope_of_work)
    _subheading(doc, "B", "Scope of Services")
    _body(doc, scope_of_services)

    # ---- VIII. Project Financials
    _heading(doc, "VIII", "Project Financials")
    _meta_line(doc, "Link to Project Set-up Form", fee_link)
    _meta_line(doc, "Invoice Frequency", invoice_frequency)
    _meta_line(doc, "Send By Date", invoice_date)

    # ---- IX. Schedule & Milestones
    _heading(doc, "IX", "Schedule & Milestones")
    _subheading(doc, "A", "Project Schedule")
    _meta_line(doc, "Link to Detailed Project Schedule", schedule_link)
    _subheading(doc, "B", "Project Milestones")
    _table(
        doc, ["#", "Milestone Description", "Date"],
        [[i + 1, m.get("description", ""), m.get("date", "")] for i, m in enumerate(milestones or [])],
        col_widths=[0.4, 4.3, 1.5],
    )
    _subheading(doc, "C", "Progress Reports")
    _meta_line(doc, "Frequency", progress_frequency)
    _meta_line(doc, "Format", progress_format)
    _meta_line(doc, "Delivery Method", progress_delivery)

    # ---- X. BIM/CAD Requirements
    _heading(doc, "X", "BIM / CAD Requirements")
    _meta_line(doc, "Link to BIM Execution Plan", bim_link)

    # ---- XI. QA and QC Plans
    _heading(doc, "XI", "QA and QC Plans")
    _body(doc, qa_qc_plan)

    # ---- XII. Risk Management Plan
    _heading(doc, "XII", "Risk Management Plan")
    _table(
        doc, ["Risk Factor", "Potential Impact", "Sev.", "Prob.", "Mitigation Measures", "By Whom", "By When"],
        [[r_.get("factor", ""), r_.get("impact", ""), r_.get("severity", ""),
          r_.get("probability", ""), r_.get("mitigation", ""), r_.get("by_whom", ""), r_.get("by_when", "")]
         for r_ in (risks or [])],
        col_widths=[1.6, 1.6, 0.5, 0.5, 2.0, 0.7, 0.9],
    )

    doc.save(str(output_path))
    return output_path
